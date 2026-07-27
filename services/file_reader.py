"""
File Reader Service - Reads content from .txt, .md, .docx, .pdf files.
Handles images by encoding to base64 for OpenAI Vision API.
Handles large files by truncating to prevent OpenAI token overflow.
"""

import os
import base64
MAX_CONTENT_CHARS = 10_000
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}

class FileReader:
    def read_file(self, filepath: str) -> str:
        """
        Read content from a file based on its extension.
        Returns plain text content, truncated if too long.
        For images, raises an error directing callers to use read_image_as_base64().
        """
        ext = os.path.splitext(filepath)[1].lower()

        if ext in IMAGE_EXTENSIONS:
            raise ValueError(
                f"File ảnh ({ext}) không thể đọc dạng text. "
                "Dùng read_image_as_base64() để lấy dữ liệu gửi lên AI Vision."
            )

        readers = {
            '.txt':  self._read_text,
            '.md':   self._read_text,
            '.docx': self._read_docx,
            '.pdf':  self._read_pdf,
        }

        reader = readers.get(ext)
        if not reader:
            raise ValueError(f'Định dạng file không hỗ trợ: {ext}')
        content = reader(filepath)
        return self._truncate(content)
    def is_image(self, filepath: str) -> bool:
        """Return True if the file is a supported image format."""
        return os.path.splitext(filepath)[1].lower() in IMAGE_EXTENSIONS
    def read_image_as_base64(self, filepath: str) -> dict:
        """
        Read an image file and return a dict ready to be embedded in an
        OpenAI Vision message content block.
        Usage in ai_service.py:
            image_block = file_reader.read_image_as_base64(filepath)
            messages = [{"role": "user", "content": [image_block, {"type": "text", "text": prompt}]}]
        """
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in IMAGE_EXTENSIONS:
            raise ValueError(f"Không phải file ảnh hợp lệ: {ext}")
        media_type_map = {
            '.jpg':  'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png':  'image/png',
            '.gif':  'image/gif',
            '.webp': 'image/webp',
            'bmp':  'image/bmp',
        }
        with open(filepath, 'rb') as f:
            image_data = base64.b64encode(f.read()).decode('utf-8')
        return {
            "type": "image_url",
            "image_url": {
                "url": f"data:{media_type_map[ext]};base64,{image_data}",
                "detail": "high",   # "high" = OCR tốt hơn; đổi "low" nếu cần tiết kiệm token
            }
        }
    def _read_text(self, filepath: str) -> str:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    def _read_docx(self, filepath: str) -> str:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx chưa được cài. Chạy: pip install python-docx")
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs)
    def _read_pdf(self, filepath: str) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError("pypdf chưa được cài. Chạy: pip install pypdf")
        reader = PdfReader(filepath)
        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
        return '\n'.join(pages)

    def _truncate(self, content: str) -> str:
        if len(content) > MAX_CONTENT_CHARS:
            return (
                content[:MAX_CONTENT_CHARS]
                + f'\n\n...[Nội dung bị cắt bớt. '
                  f'Đã đọc {MAX_CONTENT_CHARS:,} / {len(content):,} ký tự]...'
            )
        return content